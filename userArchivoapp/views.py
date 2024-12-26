from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache
from django.template.loader import render_to_string
from .forms import DocumentoForm, GestionSolicitudForm, DevolucionForm
from django.contrib import messages
from webapp.models import Secretaria, Unidad, UsuarioCargoUnidad
from userPersonal.models import SolicitudPrestamo
from PIL import Image
from django.http import JsonResponse, HttpResponse, FileResponse
import os, io
from .models import Documento, Prestamo, Gestion
from PyPDF2 import PdfReader, PdfWriter
import zipfile
from docx import Document
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import openpyxl
from io import BytesIO
from django.core.files.uploadedfile import InMemoryUploadedFile
from collections import defaultdict
from itertools import groupby
from operator import attrgetter
from django.utils import timezone
from django.utils.timezone import localtime
from django.db.models import F
from django.conf import settings
from datetime import timedelta
from operator import attrgetter
from django.core.paginator import Paginator


def comprimir_archivo(archivo):
    if archivo.name.endswith(('.jpg', '.jpeg', '.png')):
        # Comprimir imágenes
        img = Image.open(archivo)
        output = BytesIO()
        img.save(output, format='JPEG', quality=70)  # Reducir calidad para comprimir
        output.seek(0)
        return InMemoryUploadedFile(output, 'ImageField', f"{archivo.name.split('.')[0]}.jpg", 'image/jpeg', len(output.getvalue()), None)
    
    elif archivo.name.endswith('.pdf'):
        # Comprimir PDFs
        pdf_reader = PdfReader(archivo)
        pdf_writer = PdfWriter()
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)
        output = BytesIO()
        pdf_writer.write(output)
        output.seek(0)
        return InMemoryUploadedFile(output, 'FileField', archivo.name, 'application/pdf', len(output.getvalue()), None)

    elif archivo.name.endswith('.docx'):
        # Comprimir documentos Word
        doc = DocxDocument(archivo)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return InMemoryUploadedFile(output, 'FileField', archivo.name, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', len(output.getvalue()), None)

    elif archivo.name.endswith(('.xls', '.xlsx')):
        # Comprimir documentos Excel
        output = io.BytesIO()
        if archivo.name.endswith('.xlsx'):
            wb = openpyxl.load_workbook(archivo)
            wb.save(output)
        else:
            # Manejar archivos .xls si es necesario
            pass
        output.seek(0)
        return InMemoryUploadedFile(output, 'FileField', archivo.name, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', len(output.getvalue()), None)

    return archivo

@never_cache
@login_required
def subir_documento(request):
    gestion_abierta = Gestion.objects.filter(abierta=True).first()  # Obtener la gestión abierta
    form = DocumentoForm()
    
    if request.method == 'POST':
        form = DocumentoForm(request.POST, request.FILES)
        if form.is_valid() and gestion_abierta:  # Asegurarse de que hay una gestión abierta
            documento = form.save(commit=False)
            documento.usuario = request.user  # Asignar el usuario que subió el documento          
            # Convertir todos los campos relevantes a mayúsculas
            documento.nombre_archivo = form.cleaned_data['nombre_archivo'].upper()
            documento.fecha_documento = form.cleaned_data['fecha_documento']
            documento.detalles = form.cleaned_data['detalles'].upper()
            documento.estado = form.cleaned_data['estado'].upper()
            documento.numero_hojas = form.cleaned_data['numero_hojas']  # Suponiendo que este es un número
            documento.tipo_documentacion = form.cleaned_data['tipo_documentacion'].upper()
            documento.ambiente = form.cleaned_data['ambiente'].upper()
            documento.estante = form.cleaned_data['estante'].upper()
            documento.columna = form.cleaned_data['columna'].upper()
            documento.balda = form.cleaned_data['balda'].upper()
            documento.tipo = form.cleaned_data['tipo'].upper()
            documento.numero = form.cleaned_data['numero']  # Suponiendo que este es un número
            
            documento.gestion = gestion_abierta.año  # Asigna el año de la gestión abierta
            documento.responsable = form.cleaned_data['responsable'].upper()
            
            archivo = request.FILES.get('archivo')
            if archivo:
                print("Archivo recibido:", archivo.name)
                documento.archivo = comprimir_archivo(archivo)  # Suponiendo que tienes esta función
            else:
                print("No se recibió archivo.")
            documento.save()
            
            # Agregar un mensaje de éxito
            return redirect('reporte_documento', documento.id)  # Redirige a la vista del reporte
            
    unidades = Unidad.objects.all().select_related('nombre_secretaria')
    
    context = {
        'form': form,
        'unidades': unidades,
        'nombre_archivo': form['nombre_archivo'],
        'archivo': form['archivo'],
        'estado': form['estado'],
        'numero_hojas': form['numero_hojas'],
        'tipo_documentacion': form['tipo_documentacion'],
        'ambiente': form['ambiente'],
        'estante': form['estante'],
        'columna': form['columna'],
        'balda': form['balda'],
        'detalles': form['detalles'],
        'tipo': form['tipo'],
        'numero': form['numero'],
        'gestion': gestion_abierta.año if gestion_abierta else None,  # Usar el año de la gestión abierta
        'responsable': form['responsable'],
        'fecha_documento': form['fecha_documento'],
    }
    
    return render(request, 'subir_documento.html', context)

def get_unidades(request, secretaria_id):
    unidades = Unidad.objects.filter(nombre_secretaria_id=secretaria_id)
    unidades_list = list(unidades.values('id', 'nombre_unidad'))
    return JsonResponse({'unidades': unidades_list})

@login_required
@never_cache
def vista_previa_documento(request, documento_id):
    documento = get_object_or_404(Documento, id=documento_id)
    es_pdf = documento.archivo.url.lower().endswith(".pdf")
    archivo_url = request.build_absolute_uri(documento.archivo.url)
    context = {
        'documento': documento,
        'es_pdf': es_pdf,
        'archivo_url': archivo_url,
    }
    return render(request, 'vista_previa_documento.html', context)

@never_cache
@login_required
def gestionar_solicitud(request, solicitud_id):
    solicitud = get_object_or_404(SolicitudPrestamo, id=solicitud_id)
    
    if request.method == 'POST':
        form = GestionSolicitudForm(request.POST, instance=solicitud)
        if form.is_valid():
            form.save()
            return redirect('lista_solicitudes')  # Redirigir a una lista de solicitudes
    
    else:
        form = GestionSolicitudForm(instance=solicitud)
    
    return render(request, 'gestionar_solicitud.html', {'form': form, 'solicitud': solicitud})

@never_cache
@login_required
def lista_archivos(request):
    # Obtener los valores de los campos de búsqueda
    comprobante = request.GET.get('comprobante', '')
    gestion = request.GET.get('gestion', '')
    nombre = request.GET.get('nombre', '')

    # Filtrar documentos según los criterios proporcionados
    documentos = Documento.objects.all()

    if comprobante:
        documentos = documentos.filter(numero__icontains=comprobante)
    if gestion:
        documentos = documentos.filter(gestion__icontains=gestion)
    if nombre:
        documentos = documentos.filter(nombre_archivo__icontains=nombre)

    # Ordenar los documentos por 'detalles' y 'fecha_subida'
    documentos = documentos.order_by('detalles', 'fecha_subida')

    # Agrupar los documentos por 'detalles' en una lista de tuplas
    documentos_agrupados = [
        (detalle, list(docs))
        for detalle, docs in groupby(documentos, key=attrgetter('detalles'))
    ]

    # Paginador: muestra 6 grupos por página
    paginator = Paginator(documentos_agrupados, 6)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Preparar el contexto
    context = {
        'page_obj': page_obj,
        'comprobante': comprobante,
        'gestion': gestion,
        'nombre': nombre,
    }

    return render(request, 'list_doc.html', context)

#vista que maneja la logica de la lista de solicitudes
@never_cache
@login_required
def lista_solicitudes(request): 
    # Filtrar solicitudes según el estado de aprobación
    solicitudes_pendientes = SolicitudPrestamo.objects.filter(aprobado=False)
    solicitudes_aprobadas = SolicitudPrestamo.objects.filter(aprobado=True, prestamo__estado_prestamo=False)
    
    # Filtrar solicitudes rechazadas utilizando el estado en el modelo Prestamo
    solicitudes_rechazadas = SolicitudPrestamo.objects.filter(prestamo__estado_solicitud='rechazado')

    context = {
        'solicitudes_pendientes': solicitudes_pendientes,
        'solicitudes_aprobadas': solicitudes_aprobadas,
        'solicitudes_rechazadas': solicitudes_rechazadas,
    }
    return render(request, 'lista_solicitudes.html', context)


#vista de la lista de orestamos pendientes de devolocion
@never_cache
@login_required
def lista_prestamos_pendientes(request):
    # Préstamos pendientes de devolución
    prestamos_pendientes = Prestamo.objects.filter(estado_prestamo=False, estado_solicitud='aprobado')
   
    # Préstamos ya devueltos
    prestamos_devueltos = Prestamo.objects.filter(estado_prestamo=True, estado_solicitud='aprobado')
    
    # Solicitudes pendientes de aprobación
    solicitudes_pendientes = SolicitudPrestamo.objects.filter(aprobado=False)  # Filtrar por solicitudes no aprobadas

    # Pasar las listas al template
    return render(request, 'lista_prestamos_pendientes.html', {
        'prestamos_pendientes': prestamos_pendientes,
        'prestamos_devueltos': prestamos_devueltos,
        'solicitudes_pendientes': solicitudes_pendientes,  # Enviar las solicitudes pendientes
    })

@never_cache
@login_required
def registrar_devolucion_modal(request, prestamo_id):
    prestamo = get_object_or_404(Prestamo, id=prestamo_id)

    if request.method == 'POST':
        # Solo guardamos la fecha de devolución y cambiamos el estado
        prestamo.estado_prestamo = True  # Marca el préstamo como devuelto
        prestamo.fecha_devolucion = timezone.now()  # Establece la fecha de devolución actual
        prestamo.save()  # Guarda los cambios en el préstamo

        # Redirige a la página con los préstamos ya devueltos
        return redirect('lista_prestamos_pendientes')  # Cambia la URL a la que desees redirigir

    return JsonResponse({'success': False})  # Si algo sale mal, retorna un error
@never_cache
@login_required
def buscar_documentos(request):
    query = request.GET.get('query', '')
    tipo = request.GET.get('tipo', '')  # Filtro para el tipo
    fecha_subida = request.GET.get('fecha_subida', '')
    año = request.GET.get('año', '')

    resultados = Documento.objects.all()

    if query:
        resultados = resultados.filter(nombre_archivo__icontains=query)

    if tipo:
        resultados = resultados.filter(tipo__icontains=tipo)

    if fecha_subida:
        resultados = resultados.filter(fecha_subida=fecha_subida)

    if año:
        resultados = resultados.filter(fecha_subida__year=año)

    # Configurar paginador
    paginator = Paginator(resultados, 5)  # Cambia '5' al número de resultados por página
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        html = render_to_string('_resultados_busqueda.html', {'resultados': page_obj})
        return JsonResponse({'html': html})

    return render(request, 'buscar_documentos.html', {'resultados': page_obj})


@never_cache
@login_required
def reporte_documento(request, documento_id):   
    # Obtener el documento subido por su ID
    documento = get_object_or_404(Documento, id=documento_id)
    
    # Pasar el documento al contexto
    context = {
        'documento': documento
    }
    # Renderizar la plantilla del reporte
    return render(request, 'reporte_documento.html', context)

@never_cache
@login_required
def lista_reportes(request):
    filtro = request.GET.get('filtro', 'dia')
    fecha_inicio = request.GET.get('fecha_inicio', None)
    fecha_fin = request.GET.get('fecha_fin', None)
    documentos = Documento.objects.all()

    # Filtrado por fecha
    if filtro == 'dia':
        hoy = localtime(timezone.now()).date()
        documentos = documentos.filter(fecha_subida__date=hoy)
        print(f"Filtrando documentos del día: {hoy}, encontrados: {documentos.count()}")
    elif filtro == 'semana':
        semana_pasada = timezone.now() - timedelta(weeks=1)
        documentos = documentos.filter(fecha_subida__gte=semana_pasada)
    elif filtro == 'mes':
        mes_pasado = timezone.now() - timedelta(days=30)
        documentos = documentos.filter(fecha_subida__gte=mes_pasado)
    elif filtro == 'rango' and fecha_inicio and fecha_fin:
        try:
            fecha_inicio = timezone.datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
            fecha_fin = timezone.datetime.strptime(fecha_fin, '%Y-%m-%d').date()
            documentos = documentos.filter(fecha_subida__date__range=(fecha_inicio, fecha_fin))
        except ValueError:
            print("Error en el rango de fechas proporcionado.")
            documentos = documentos

    # Paginación
    paginator = Paginator(documentos, 10)  # 10 documentos por página
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'lista_reportes.html', {
        'documentos': page_obj,
        'filtro': filtro
    })

@login_required
def descargar_reporte(request, documento_id):
    # Obtener el documento por ID
    documento = get_object_or_404(Documento, id=documento_id)
   
    # Crear un nuevo documento de Word
    doc = Document()

    # Ruta de la imagen del encabezado
    imagen_encabezado_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'ESCUDO TEXTO NEGRO HORIZONTAL.png')

    # Insertar el encabezado con imagen
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    if os.path.exists(imagen_encabezado_path):
        run.add_picture(imagen_encabezado_path, width=Inches(2.5))  # Ajusta el tamaño según necesites
    header_paragraph.alignment = 1  # Centrar encabezado

    # Añadir un título al documento
    doc.add_heading('Reporte del Documento Subido', level=1)

    # Crear la tabla
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'  # Aplicar un estilo de tabla

    # Agregar filas a la tabla
    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    # Llenar la tabla con la información del documento
    add_row("Nombre del archivo:", documento.nombre_archivo)
    add_row("Detalles:", documento.detalles)
    add_row("Estado:", documento.estado)
    add_row("Número de Hojas:", str(documento.numero_hojas))
    add_row("Tipo de Documentación:", documento.tipo_documentacion)
    add_row("Ambiente:", documento.ambiente)
    add_row("Estante:", documento.estante)
    add_row("Columna:", documento.columna)
    add_row("Balda:", documento.balda)
    add_row("Tipo:", documento.tipo)
    add_row("Número:", documento.numero)
    add_row("Gestión:", documento.gestion)
    add_row("Responsable:", documento.responsable)

    # Agregar espacio para firmas al final del documento
    doc.add_paragraph("\n\n")
    firmas_table = doc.add_table(rows=1, cols=3)
    firmas_table.style = 'Table Grid'

    # Configurar las celdas de la tabla de firmas
    firmas_cells = firmas_table.rows[0].cells
    firmas_cells[0].text = "Firma del Usuario\n(Nombre del Usuario)"
    firmas_cells[1].text = "Responsable de Archivos"

    # Configuración opcional: Centrar el texto en las celdas de la tabla de firmas
    for cell in firmas_cells:
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.alignment = 1  # Centrar texto (1 = Center)

    # Configurar la respuesta HTTP para la descarga
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{documento.nombre_archivo}.docx"'

    # Guardar el documento en la respuesta
    doc.save(response)

    return response

def descargar_documento(request, documento_id):
    # Obtener el documento por su ID
    documento = get_object_or_404(Documento, id=documento_id)

    # Abrir el archivo en modo lectura binaria
    file_path = documento.archivo.path  # Supongo que tu modelo Documento tiene un campo llamado archivo
    response = FileResponse(open(file_path, 'rb'), as_attachment=True)

    # Asegurarte de que el nombre del archivo tenga la extensión .pdf
    nombre_archivo = documento.nombre_archivo
    if not nombre_archivo.lower().endswith('.pdf'):
        nombre_archivo += '.pdf'

    # Establecer el nombre del archivo para la descarga
    response['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'

    return response

@never_cache
@login_required
def gestionar_gestion(request):
    if request.method == 'POST':
        año = request.POST.get('año')
        Gestion.objects.update(abierta=False)  # Cierra cualquier gestión anterior
        gestion, created = Gestion.objects.get_or_create(año=año)
        gestion.abierta = True
        gestion.save()
        return redirect('gestionar_gestion')

    gestiones = Gestion.objects.all()
    return render(request, 'gestionar_gestion.html', {'gestiones': gestiones})
def cerrar_gestion(request, gestion_id):
    gestion = get_object_or_404(Gestion, id=gestion_id)
    gestion.abierta = False
    gestion.save()
    return redirect('gestionar_gestion')