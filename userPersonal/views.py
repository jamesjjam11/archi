from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.views.decorators.cache import never_cache
from django.http import HttpResponse
from .forms import SolicitudPrestamoForm
from .models import SolicitudPrestamo
from webapp.models import Persona
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

@never_cache
@login_required
def solicitar_prestamo_view(request):
    if request.method == 'POST':  
        form = SolicitudPrestamoForm(request.POST, user=request.user)
        if form.is_valid():
            # Guarda la solicitud en la base de datos
            solicitud = form.save(commit=False)
            solicitud.usuario = request.user
            solicitud.aprobado = False  # Solicitud pendiente
            solicitud.save()

            # Obtener datos del modelo Persona
            persona = Persona.objects.get(user=request.user)

            # Genera el documento Word
            buffer = BytesIO()
            doc = Document()

            # Añadir encabezado con imagen
            header = doc.sections[0].header
            header_paragraph = header.paragraphs[0]
            header_run = header_paragraph.add_run()
            header_run.add_picture('static/img/ESCUDO TEXTO NEGRO HORIZONTAL.png', width=Inches(2.0))
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Añadir título
            doc.add_heading('Solicitud de Préstamo de Documentos', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Crear tabla para organizar datos
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Table Grid'
            data = [
                ("Nombre", persona.nombre),
                ("Apellido Paterno", persona.apellido_paterno),
                ("Apellido Materno", persona.apellido_materno),
                ("Cargo", form.cleaned_data['cargo']),
                ("Secretaría", form.cleaned_data['secretaria']),
                ("Unidad", form.cleaned_data['unidad']),
                ("Fecha de Solicitud", solicitud.fecha_solicitud.strftime('%d/%m/%Y')),
            ]
            for i, (key, value) in enumerate(data):
                table.cell(i, 0).text = key
                table.cell(i, 1).text = str(value)

            # Añadir motivo de la solicitud
            doc.add_paragraph("\nMotivo de la Solicitud:")
            doc.add_paragraph(form.cleaned_data['motivo_solicitud'], style='Intense Quote')

            # Añadir información adicional
            doc.add_paragraph("\nDocumento Solicitado:")
            doc.add_paragraph(f"{solicitud.documento}", style='Intense Quote')

            # Añadir espacios para las firmas al final
            doc.add_paragraph("\n\n")  # Espacio adicional para separar contenido previo

            # Añadir una tabla con espacios para firmas
            firma_table = doc.add_table(rows=2, cols=3)
            firma_table.style = 'Table Grid'

            # Primera fila: Títulos de las firmas
            titles = [
                "Firma del Interesado\n(Nombre del Usuario)",
                "Responsable de la Unidad de Archivos",
                "Responsable de la Unidad del Solicitante",
            ]
            for i, title in enumerate(titles):
                cell = firma_table.cell(0, i)
                cell.text = title
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(10)

            # Segunda fila: Espacios para las firmas
            for i in range(3):
                cell = firma_table.cell(1, i)
                paragraph = cell.paragraphs[0]
                paragraph.add_run("\n\n\n_________________________").font.size = Pt(10)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Guardar el documento en el buffer
            doc.save(buffer)
            buffer.seek(0)

            # Enviar el archivo como respuesta HTTP
            response = HttpResponse(
                buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="Solicitud_Prestamo.docx"'
            return response
    else:
        form = SolicitudPrestamoForm(user=request.user)
    return render(request, 'solicitud_prestamo.html', {'form': form})
